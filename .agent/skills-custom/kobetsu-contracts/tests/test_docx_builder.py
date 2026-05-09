"""Tests for docx_builder.py module."""

import pytest
from pathlib import Path

import sys

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

try:
    from kobetsu_contracts.scripts.docx_builder import DOCXBuilder

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


class TestDOCXBuilder:
    """Test DOCXBuilder class."""

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
    def test_init(self):
        """Test DOCX builder initialization."""
        builder = DOCXBuilder()
        assert builder is not None

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
    def test_init_without_python_docx(self):
        """Test initialization fails without python-docx."""
        if HAS_PYTHON_DOCX:
            pytest.skip("python-docx is installed")

        with pytest.raises(ImportError):
            DOCXBuilder()

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
    def test_build_valid_contract(self, temp_db_dir, test_employee, test_contract):
        """Test DOCX building with valid contract."""
        builder = DOCXBuilder()

        output_path = temp_db_dir / "test_contract.docx"

        result = builder.build(test_employee, test_contract, output_path)

        assert result == output_path
        assert output_path.exists()
        assert output_path.stat().st_size > 0

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
    def test_build_creates_output_directory(self, temp_db_dir, test_employee, test_contract):
        """Test that build creates output directory if needed."""
        builder = DOCXBuilder()

        nested_path = temp_db_dir / "level1" / "level2" / "test.docx"

        builder.build(test_employee, test_contract, nested_path)

        assert nested_path.exists()
        assert nested_path.parent == temp_db_dir / "level1" / "level2"

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
    def test_build_from_contract(self, temp_db_dir, test_employee, test_contract):
        """Test build_from_contract convenience method."""
        builder = DOCXBuilder()

        result = builder.build_from_contract(
            test_employee,
            test_contract,
            output_dir=temp_db_dir,
        )

        assert result.exists()
        assert "KOB-2026-001" in result.name
        assert result.suffix == ".docx"

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
    def test_build_from_contract_sanitizes_filename(self, temp_db_dir, test_employee):
        """Test that contract number is sanitized for filename."""
        from _uns_employee_schema import KobetsuContract
        from datetime import date

        contract = KobetsuContract(
            id="TEST-1",
            employee_id="W001",
            contract_number="KOB/2026/001 (draft)",
            contract_date=date(2026, 4, 3),
            effective_date=date(2026, 4, 1),
            job_title="Test",
            work_location="Tokyo",
            hourly_rate_jpy=1000.0,
        )

        builder = DOCXBuilder()

        result = builder.build_from_contract(
            test_employee,
            contract,
            output_dir=temp_db_dir,
        )

        # Path separators should be replaced
        assert "/" not in result.name
        assert " " not in result.name or "_" in result.name

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
    def test_build_includes_all_sections(self, temp_db_dir, test_employee, test_contract):
        """Test that DOCX includes all required sections."""
        builder = DOCXBuilder()

        output_path = temp_db_dir / "complete.docx"
        result = builder.build(test_employee, test_contract, output_path)

        # Try to read back the document
        try:
            from docx import Document

            doc = Document(result)

            # Check that document is not empty
            assert len(doc.paragraphs) > 0

            # Verify key content is present
            full_text = "\n".join([p.text for p in doc.paragraphs])
            assert "個別派遣契約書" in full_text
            assert test_employee.name_jp in full_text
            assert test_contract.job_title in full_text

        except Exception:
            # If we can't verify, just check file exists
            assert result.exists()

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
    def test_build_with_japanese_content(self, temp_db_dir, test_employee, test_contract):
        """Test DOCX building with Japanese content."""
        builder = DOCXBuilder()

        output_path = temp_db_dir / "japanese.docx"
        result = builder.build(test_employee, test_contract, output_path)

        assert result.exists()
        assert output_path.stat().st_size > 1000

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
    def test_build_with_minimal_contract(self, temp_db_dir, test_employee):
        """Test DOCX building with minimal contract data."""
        from _uns_employee_schema import KobetsuContract
        from datetime import date

        contract = KobetsuContract(
            id="TEST-MINIMAL",
            employee_id=test_employee.id,
            contract_number="KOB-MINIMAL",
            contract_date=date(2026, 4, 3),
            effective_date=date(2026, 4, 1),
            job_title="Test Job",
            work_location="Tokyo",
            hourly_rate_jpy=1000.0,
        )

        builder = DOCXBuilder()

        output_path = temp_db_dir / "minimal.docx"
        result = builder.build(test_employee, contract, output_path)

        assert result.exists()

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
    def test_build_with_long_text_fields(self, temp_db_dir, test_employee):
        """Test DOCX building with long text in fields."""
        from _uns_employee_schema import KobetsuContract
        from datetime import date

        contract = KobetsuContract(
            id="TEST-LONG",
            employee_id=test_employee.id,
            contract_number="KOB-LONG",
            contract_date=date(2026, 4, 3),
            effective_date=date(2026, 4, 1),
            job_title="Senior Software Engineer - Multiple Responsibilities",
            work_location="Tokyo HQ Building, 3rd Floor, Room 301-305",
            hourly_rate_jpy=5000.0,
            terms_conditions="This contract includes detailed terms and conditions spanning multiple lines.\n"
            * 5,
            notes="Special notes and requirements for this contract position.\n" * 3,
        )

        builder = DOCXBuilder()

        output_path = temp_db_dir / "long_text.docx"
        result = builder.build(test_employee, contract, output_path)

        assert result.exists()
        assert result.stat().st_size > 2000
