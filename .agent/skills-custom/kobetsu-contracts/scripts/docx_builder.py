"""DOCX Builder for Kobetsu contracts using python-docx.

Generates Microsoft Word documents (.docx) for contracts with professional formatting.
"""

import logging
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

import sys

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from _uns_employee_schema import KobetsuContract, Employee

logger = logging.getLogger(__name__)


class DOCXBuilder:
    """Builds DOCX (Microsoft Word) documents from contract data."""

    def __init__(self) -> None:
        """Initialize DOCX builder.

        Raises:
            ImportError: If python-docx is not installed.
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX generation. Install with: pip install python-docx"
            )
        logger.info("DOCXBuilder initialized with python-docx")

    def build(
        self,
        employee: Employee,
        contract: KobetsuContract,
        output_path: Path,
    ) -> Path:
        """Build DOCX contract document.

        Args:
            employee: Employee record.
            contract: Contract record.
            output_path: Path where DOCX will be saved.

        Returns:
            Path to generated DOCX file.

        Raises:
            ValueError: If employee or contract data is invalid.
            Exception: If DOCX generation fails.
        """
        try:
            # Create output directory if needed
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Create document
            doc = Document()

            # Title
            title = doc.add_paragraph()
            title_run = title.add_run("個別派遣契約書")
            title_run.font.size = Pt(20)
            title_run.font.bold = True
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Subtitle
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run("(Individual Dispatch Contract)")
            subtitle_run.font.size = Pt(12)
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add spacing
            doc.add_paragraph()

            # Contract header information
            self._add_section(
                doc,
                "契約情報",
                [
                    ("契約番号", contract.contract_number),
                    ("契約日", contract.contract_date.isoformat()),
                    (
                        "有効期間",
                        f"{contract.effective_date.isoformat()} ～ {contract.expiration_date.isoformat() if contract.expiration_date else '未定'}",
                    ),
                ],
            )

            # Employee information
            self._add_section(
                doc,
                "被派遣者情報",
                [
                    ("氏名", employee.name_jp),
                    ("従業員ID", employee.id),
                    ("メールアドレス", employee.email),
                    ("電話番号", employee.phone or ""),
                ],
            )

            # Work details
            self._add_section(
                doc,
                "勤務条件",
                [
                    ("派遣会社", contract.dispatch_company),
                    ("勤務先企業", contract.client_company or ""),
                    ("勤務地", contract.work_location),
                    ("職務内容", contract.job_title),
                    ("日当たり勤務時間", f"{contract.work_hours_daily}時間"),
                    ("週間勤務日数", f"{contract.work_days_weekly}日"),
                    ("時給", f"¥{contract.hourly_rate_jpy:,.0f}"),
                    (
                        "月間最小勤務時間",
                        str(contract.monthly_minimum_hours)
                        if contract.monthly_minimum_hours
                        else "未定",
                    ),
                ],
            )

            # Benefits
            if contract.benefits:
                benefits_text = "\n".join([f"• {b}" for b in contract.benefits])
                self._add_section(
                    doc,
                    "福利厚生",
                    [
                        ("内容", benefits_text),
                    ],
                )

            # Terms and conditions
            if contract.terms_conditions:
                self._add_section(
                    doc,
                    "契約条件",
                    [
                        ("詳細", contract.terms_conditions),
                    ],
                )

            # Notes
            if contract.notes:
                self._add_section(
                    doc,
                    "備考",
                    [
                        ("内容", contract.notes),
                    ],
                )

            # Signature section
            doc.add_paragraph()
            doc.add_paragraph()

            sig_table = doc.add_table(rows=3, cols=2)
            sig_table.autofit = False

            # Row 1: Dispatch company
            cell_label = sig_table.rows[0].cells[0]
            cell_label.text = "派遣会社"
            sig_table.rows[0].cells[1].text = "UNS"

            # Row 2: Company representative
            cell_label = sig_table.rows[1].cells[0]
            cell_label.text = "代表者署名"
            sig_table.rows[1].cells[1].text = "_________________"

            # Row 3: Date
            cell_label = sig_table.rows[2].cells[0]
            cell_label.text = "日付"
            sig_table.rows[2].cells[1].text = f"{contract.contract_date.isoformat()}"

            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer_run = footer.add_run(
                f"生成: {contract.generated_at} | テンプレートバージョン: {contract.template_version}"
            )
            footer_run.font.size = Pt(8)
            footer_run.font.italic = True
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Save document
            doc.save(str(output_path))

            logger.info(f"Generated DOCX: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            raise

    def _add_section(
        self,
        doc: Document,
        section_title: str,
        fields: list,
    ) -> None:
        """Add a formatted section to the document.

        Args:
            doc: Document object.
            section_title: Section header text.
            fields: List of tuples (label, value).
        """
        # Section header
        header = doc.add_paragraph()
        header_run = header.add_run(section_title)
        header_run.font.size = Pt(12)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 51, 102)

        # Fields table
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Light Grid Accent 1"

        for idx, (label, value) in enumerate(fields):
            row = table.rows[idx]
            row.cells[0].text = label
            row.cells[1].text = str(value)

            # Format label cell
            label_paragraph = row.cells[0].paragraphs[0]
            label_run = label_paragraph.runs[0]
            label_run.font.bold = True

        doc.add_paragraph()  # Spacing

    def build_from_contract(
        self,
        employee: Employee,
        contract: KobetsuContract,
        output_dir: Path | None = None,
    ) -> Path:
        """Build DOCX using contract number in filename.

        Args:
            employee: Employee record.
            contract: Contract record.
            output_dir: Output directory (default: ./output/docx/).

        Returns:
            Path to generated DOCX file.
        """
        if output_dir is None:
            output_dir = Path.cwd() / "output" / "docx"

        # Sanitize contract number for filename
        safe_filename = contract.contract_number.replace("/", "_").replace(" ", "_")
        output_path = output_dir / f"{safe_filename}.docx"

        return self.build(
            employee=employee,
            contract=contract,
            output_path=output_path,
        )
